import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def setup_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # 1.15 line spacing roughly equals 1.15 in docx terms
    style.paragraph_format.line_spacing = 1.15
    return doc

def create_proposal(filepath):
    doc = setup_document()

    # COVER PAGE
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("AI-Driven Security Operations with Blockchain Accountability & Insurance Verification\n")
    title_run.font.size = Pt(24)
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run("Final Implementation Report\n\n")
    sub_run.font.size = Pt(18)
    
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    auth_run = authors.add_run("Authors: Omar Dorgham, Mariam El-Kholey\nCourse: AID 325 - Blockchain & Distributed Ledgers\nDate: May 16, 2026")
    auth_run.font.size = Pt(14)
    
    doc.add_page_break()

    # EXECUTIVE SUMMARY
    h = doc.add_heading("2. EXECUTIVE SUMMARY", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("The problem: AI accountability + insurance fraud.")
    doc.add_paragraph("The solution: Blockchain-anchored audit trail + automated claims.")
    doc.add_paragraph("Key innovation: Integration of AI defensive actions into insurance payout logic.")
    doc.add_page_break()

    # SYSTEM ARCHITECTURE
    h = doc.add_heading("3. SYSTEM ARCHITECTURE", level=1)
    h.style.font.size = Pt(14)
    doc.add_heading("AI Accountability Module", level=2)
    doc.add_paragraph("• Off-chain: Full event database with reasoning text\n• On-chain: Merkle roots every 60 seconds\n• Daily reports: Summary hashes submitted to DailyReportRegistry", style='List Bullet')
    
    doc.add_heading("Cyber Insurance Module", level=2)
    doc.add_paragraph("• Daily posture snapshots (PostureRegistry)\n• Policy rules (PolicyEngine)\n• Automated claims (ClaimsProcessor)", style='List Bullet')
    
    doc.add_heading("INTEGRATION POINT", level=2)
    doc.add_paragraph("ClaimsProcessor queries AuditRegistry to verify AI defensive actions.")
    doc.add_page_break()

    # SMART CONTRACTS
    h = doc.add_heading("4. SMART CONTRACTS", level=1)
    h.style.font.size = Pt(14)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Contract'
    hdr_cells[1].text = 'Purpose'
    hdr_cells[2].text = 'Key Functions'
    
    contracts = [
        ("AuditRegistry", "AI action logging", "logAction(), getActionCountByTimeRange()"),
        ("Governance", "Multi-sig approval", "approveAction()"),
        ("PostureRegistry", "Daily security snapshots", "recordPosture()"),
        ("PolicyEngine", "Coverage rules", "checkCompliance()"),
        ("ClaimsProcessor", "Automated payout", "processClaim()"),
        ("CyberToken (ERC20)", "Claim payments", "transfer(), balanceOf()"),
        ("DailyReportRegistry", "Daily summaries", "submitDailyReport()")
    ]
    for contract, purpose, funcs in contracts:
        row_cells = table.add_row().cells
        row_cells[0].text = contract
        row_cells[1].text = purpose
        row_cells[2].text = funcs
        
    doc.add_page_break()

    # TECHNICAL IMPLEMENTATION
    h = doc.add_heading("5. TECHNICAL IMPLEMENTATION", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("• Tech Stack: Hardhat, Solidity 0.8.x, Python FastAPI, React.js, Gemini AI", style='List Bullet')
    doc.add_paragraph("• Security Measures: ReentrancyGuard, CEI pattern, access control modifiers", style='List Bullet')
    doc.add_paragraph("• Testing: 4 unit tests covering happy path, edge cases, and security", style='List Bullet')
    doc.add_page_break()

    # INTEGRATION LOGIC
    h = doc.add_heading("6. INTEGRATION LOGIC", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("When a breach occurs and a claim is filed, ClaimsProcessor:")
    doc.add_paragraph("1. Checks PostureRegistry for pre-breach security compliance\n2. Queries AuditRegistry for AI defensive actions in the 7 days before breach\n3. If AI took ≥5 HIGH-risk defensive actions, compliance score increases 10%\n4. Final payout = baseAmount × (complianceScore / 100)", style='List Number')
    doc.add_page_break()

    # FUTURE WORK
    h = doc.add_heading("7. FUTURE WORK", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("• Deploy to Sepolia testnet\n• Integrate with real SOC alert systems\n• Add machine learning for fraud detection", style='List Bullet')
    doc.add_page_break()

    # CONCLUSION
    h = doc.add_heading("8. CONCLUSION", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("• Summary of deliverables: End-to-end integration of AI accountability and automated blockchain insurance.\n• Regulatory compliance: Aligned with EU AI Act Article 12.\n• Market impact: Targeting the $16B cyber insurance industry.", style='List Bullet')

    doc.save(filepath)

def create_threat_model(filepath):
    doc = setup_document()

    # TITLE
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Security Threat Model\n")
    title_run.font.size = Pt(18)
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run("AI-Driven Security + Blockchain Insurance System\n")
    sub_run.font.size = Pt(14)
    
    # SECTION 1
    h = doc.add_heading("SECTION 1: THREAT IDENTIFICATION (Top 5 Risks)", level=1)
    h.style.font.size = Pt(14)
    
    threats = [
        ("1. Reentrancy Attack on ClaimsProcessor", "Vector: Malicious contract calls processClaim(), recursively drains funds during token transfer\nMitigation: OpenZeppelin ReentrancyGuard + CEI pattern implemented\nResidual Risk: Low"),
        ("2. Flash Loan Governance Attack", "Vector: Attacker borrows large amount of CyberToken, gains voting power, manipulates claim decision\nMitigation: NOT YET IMPLEMENTED — requires time-locked voting or snapshot-based governance\nResidual Risk: High"),
        ("3. Timestamp Manipulation", "Vector: Miner manipulates block.timestamp to submit fraudulent daily reports or posture snapshots\nMitigation: Use block.number instead of block.timestamp for time-sensitive logic\nResidual Risk: Medium"),
        ("4. Off-Chain Database Compromise", "Vector: Attacker gains access to backend database, modifies AI reasoning text before Merkle root submission\nMitigation: Append-only log architecture, cryptographic hashing, regular backups\nResidual Risk: Medium"),
        ("5. Oracle Failure (if external price feeds added)", "Vector: Chainlink or other oracle provides stale/incorrect data for claim amounts\nMitigation: NOT APPLICABLE (system currently uses fixed token amounts)\nResidual Risk: N/A")
    ]
    
    for t_title, t_desc in threats:
        doc.add_heading(t_title, level=2)
        doc.add_paragraph(t_desc)

    # SECTION 2
    h = doc.add_heading("SECTION 2: CRITICAL VULNERABILITY ASSESSMENT", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("Current Security Posture: MEDIUM\n✅ Reentrancy protection implemented\n✅ Access control modifiers enforced\n❌ No flash loan protection\n❌ No formal audit completed")

    # SECTION 3
    h = doc.add_heading("SECTION 3: RECOMMENDATIONS FOR PRODUCTION", level=1)
    h.style.font.size = Pt(14)
    doc.add_paragraph("1. Add time-lock to governance voting (7-day delay)\n2. Implement emergency pause function (OpenZeppelin Pausable)\n3. Conduct third-party security audit\n4. Add circuit breakers for abnormal payout amounts\n5. Implement rate limiting on claim submissions", style='List Number')

    doc.save(filepath)

if __name__ == "__main__":
    proposal_path = r"c:\Dev\blockchain-project\Hybrid_Blockchain_Project\Updated_Proposal_May2026.docx"
    threat_model_path = r"c:\Dev\blockchain-project\Hybrid_Blockchain_Project\Threat_Model_Analysis.docx"
    
    print("Generating Proposal Document...")
    create_proposal(proposal_path)
    print("Generating Threat Model Document...")
    create_threat_model(threat_model_path)
    print("Generation Complete.")
